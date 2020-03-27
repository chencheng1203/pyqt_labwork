"""
根据不同的后缀名进行文本内容识别
"""
import docx
import os
import re
import fitz
from koutu import watershed,fenge
from cnocr import CnOcr
import cv2
import shutil


def find_keywords(text_path,keyword,number=2):
    """
    text_path:文本路径
    number：关键词左右浮动距离
    keyword：关键词
    """
    textlist = []
    with open(text_path, encoding='utf-8') as f:
        for line in f:
            textlist.append(line)
    
    count = 0
    txt_result = []
    for sentence in textlist:
        count += 1         #用来算行数
        flag = re.search(keyword,str(sentence))
        if flag:
            position = re.search(keyword,str(sentence)).span()
            #print("第{}行出现关键词".format(count))
            temp = "line:{}:".format(count)
            #在中间
            if position[0]>=int(number) and position[1]+int(number)<=len(sentence):
                result = sentence[int(position[0])-int(number):int(position[1])+int(number)]
                txt_result.append(temp+result)         
                #print(sentence[int(position[0])-int(number):int(position[1])+int(number)])
            #左侧出头
            elif position[0]<int(number) and position[1]+int(number)< len(sentence):
                result = sentence[:int(position[1])+int(number)]
                txt_result.append(temp+result)
                #print(sentence[:int(position[1])+int(number)])
            #右侧出头
            elif position[0]>=int(number) and position[1]+int(number)>len(sentence):
                result = sentence[position[0]+int(number):]
                txt_result.append(temp+result)
                #print(sentence[position[0]+int(number):])
            else:
                result = sentence[:]
                txt_result.append(temp+result)
                #print(sentence[:])
    return txt_result

def get_file_list():
    root = os.getcwd()
    file_list = os.listdir()
    new_list = []
    for i in range(len(file_list)):
        new_list.append(root+'/'+file_list[i])
    return new_list

def find_keyword_position(sentence,keyword):
    p = []
    f = re.search(keyword,str(sentence))
    if f == None:
        return None
    else:
        while True:
            f = re.search(keyword,str(sentence))
            if f == None:
                break
            temp = re.search(keyword,str(sentence)).span()
            p.append(temp[1])

            sentence = sentence[temp[1]:]
        
        true_p = [p[0]]
        if len(p)>1:                                  #保证多个关键词，未测试
            for i in range(1,len(p)):
                true_p.append(true_p[i-1]+p[i])
        return true_p

def find_keyword_context(sentence,keyword,true_p,number=2):
    keyword_len = len(keyword)
    sentence_len = len(sentence)
    result = []
    for i in range(len(true_p)):
        
        start = true_p[i]-keyword_len
        end = true_p[i]
        if start - number >= -1 and end +number<= sentence_len-1: 
            context = sentence[start-number:end+number]
            #print(context)
            result.append(context)
        elif start -number < -1 and end+number <= sentence_len-1:
            context = sentence[0:end+number]
            #print(context)
            result.append(context)
        elif start - number >= -1 and end +number > sentence_len-1:
            context = sentence[start-number:]
            #print(context)
            result.append(context)
        else:
            context = sentence
            #print(context)
            result.append(context)
    return result

def keyword4word(path,keyword,number):
    """
    有关键词返回关键词上下文列表
    否则换回空列表
    """
    total_result = []
    file = docx.Document(path)
    file_paragraphs_num = len(file.paragraphs)
    for i in range(file_paragraphs_num):
        sentence = file.paragraphs[i].text
        true_p = find_keyword_position(file.paragraphs[i].text, keyword)
        if true_p != None:
            result = find_keyword_context(file.paragraphs[i].text, keyword = keyword, true_p=true_p)
        else:
            continue
        for one_line in result:
            total_result.append(one_line)
    return total_result

def pdf2pic(path):
    """
    这里pdf生成完图片后先进行分割，处理完后的图片
    再送入检测
    """
    #path = '1.pdf'
    doc = fitz.open(path)
    print(path)
    #print(type(doc))
    rotate = int(0)  # 设置图片的旋转角度
    zoom_x = 2.0  # 设置图片相对于PDF文件在X轴上的缩放比例
    zoom_y = 2.0  # 设置图片相对于PDF文件在Y轴上的缩放比例
    trans = fitz.Matrix(zoom_x, zoom_y).preRotate(rotate)
    pic_dir =  './pdf2picture/'
    if not os.path.exists(pic_dir):
        os.mkdir(pic_dir)
    for pg in range(doc.pageCount):
        page = doc[pg]
        pm = page.getPixmap(matrix=trans,alpha= False)
        #print(type(pm))
        new_name = path.split("/")[-1].split('.')[0]
        #print(os.path.join(pic_dir,"%s_%s.jpg"%(new_name,pg)))
        #print(pic_dir)
        pm.writeImage(os.path.join(pic_dir,"%s_%s.jpg"%(new_name,pg)))
    
    return pic_dir

def pic_ocr(file_path,ocr,keyword,number):
    #ocr = CnOcr()
    res = ocr.ocr(file_path)
    #print(res)
    #line = []
    count = 0 
    txt_result = []
    for i in range(len(res)):
        temp_line = ''.join(res[i])
        #line.append(temp_line)
        #for sentence in temp_line:
        sentence = temp_line
        count += 1
        flag = re.search(keyword,str(sentence))
        if flag:
            # position = re.search(keyword,str(sentence)).span()
            # print("第{}行出现关键词".format(count))
            # #在中间
            # if position[0]>=int(number) and position[1]+int(number)<=len(sentence):           
            #     print(sentence[int(position[0])-int(number):int(position[1])+int(number)])
            # #左侧出头
            # elif position[0]<int(number) and position[1]+int(number)< len(sentence):
            #     print(sentence[:int(position[1])+int(number)])
            # #右侧出头
            # elif position[0]>=int(number) and position[1]+int(number)>len(sentence):
            #     print(sentence[position[0]+int(number):])
            # else:
            #     print(sentence[:])
            position = re.search(keyword,str(sentence)).span()
            #print("第{}行出现关键词".format(count))
            temp = "line:{}:".format(count)
            #在中间
            if position[0]>=int(number) and position[1]+int(number)<=len(sentence):
                result = sentence[int(position[0])-int(number):int(position[1])+int(number)]
                txt_result.append(temp+result)         
                #print(sentence[int(position[0])-int(number):int(position[1])+int(number)])
            #左侧出头
            elif position[0]<int(number) and position[1]+int(number)< len(sentence):
                result = sentence[:int(position[1])+int(number)]
                txt_result.append(temp+result)
                #print(sentence[:int(position[1])+int(number)])
            #右侧出头
            elif position[0]>=int(number) and position[1]+int(number)>len(sentence):
                result = sentence[position[0]+int(number):]
                txt_result.append(temp+result)
                #print(sentence[position[0]+int(number):])
            else:
                result = sentence[:]
                txt_result.append(temp+result)
                #print(sentence[:])
    #text_str = ''.join(line)
    return txt_result


def get_file_keyword_all(keyword, number):
    ocr = CnOcr()
    root = os.getcwd()
    #print(root)
    fileList = get_file_list()
    file_dic = {}
    file_keyword = {}

    for f in fileList:
        if f.endswith('.pdf'):
            pdf_result = []
            if os.path.exists('./pdf2picture'):
                shutil.rmtree('./pdf2picture')
            if os.path.exists('./ocr_result'):
                shutil.rmtree('./ocr_result')

            pic_dir = pdf2pic(f)               #将pdf转成图片
            file_list = os.listdir(pic_dir)    #某个pdf的所有图片
            os.mkdir('./ocr_result')
            for i in file_list:
                if i.endswith('.jpg'):
                    img_path = './pdf2picture/' + i
                    #print(img_path)
                    try:
                        img = fenge(img_path)
                        cv2.imwrite('./ocr_result/'+i, img)
                    except ValueError:
                        txt_result = pic_ocr(img_path, ocr, keyword=keyword, number=number)
                        if txt_result != []:
                            for one_line in txt_result:
                                pdf_result.append(one_line)
            ocr_list = os.listdir('./ocr_result')
            for i in ocr_list:
                i = './ocr_result/'+i
                #print(i)
                txt_result = pic_ocr(i,ocr,keyword=keyword,number =number)
                #print(txt_result)
                if txt_result != []:
                    for one_line in txt_result:
                        pdf_result.append(one_line)
            
            if pdf_result != []:
                file_dic[f] = pdf_result
                file_keyword[f] = keyword

        elif f.endswith('.txt'):
            txt_result = find_keywords(f,keyword=keyword,number=number)
            if txt_result != []:
                file_dic[f] = txt_result
                file_keyword[f] = keyword

        elif f.endswith('.docx'):
            total_result = keyword4word(f,keyword=keyword,number=number)
            if total_result != []:
                file_dic[f] = total_result
                file_keyword[f] = keyword
        else:
            continue

    count_data = file_dic.values()
    count = 0
    for i in count_data:
        for j in i:
            count += 1
    return file_dic, count


if __name__ == "__main__":
    datas = get_file_keyword_all("坦克", 2)[0]
    print(datas)