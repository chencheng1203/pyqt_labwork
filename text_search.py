# -*- encoding:utf-8 -*-
# @Time: 2020/3/19 11:48
"""
根据不同的后缀名进行文本内容识别
"""
import docx
import os
import re


def find_keywords(text_path, keyword, number=2):
    """
    text_path:文本路径
    number：关键词左右浮动距离
    keyword：关键词
    """
    textlist = []
    with open(text_path) as f:
        for line in f:
            textlist.append(line)

    count = 0
    txt_result = []
    for sentence in textlist:
        count += 1
        flag = re.search(keyword, str(sentence))
        if flag:
            position = re.search(keyword, str(sentence)).span()
            # print("第{}行出现关键词".format(count))
            temp = "第{}行出现关键词:".format(count)
            # 在中间
            if position[0] >= int(number) and position[1] + int(number) <= len(sentence):
                result = sentence[int(position[0]) - int(number):int(position[1]) + int(number)]
                txt_result.append(temp + result)
                # print(sentence[int(position[0])-int(number):int(position[1])+int(number)])
            # 左侧出头
            elif position[0] < int(number) and position[1] + int(number) < len(sentence):
                result = sentence[:int(position[1]) + int(number)]
                txt_result.append(temp + result)
                # print(sentence[:int(position[1])+int(number)])
            # 右侧出头
            elif position[0] >= int(number) and position[1] + int(number) > len(sentence):
                result = sentence[position[0] + int(number):]
                txt_result.append(temp + result)
                # print(sentence[position[0]+int(number):])
            else:
                result = sentence[:]
                txt_result.append(temp + result)
                # print(sentence[:])
    return txt_result


def get_file_list():
    root = os.getcwd()
    file_list = os.listdir()
    new_list = []
    for i in range(len(file_list)):
        new_list.append(root + '/' + file_list[i])
    return new_list


def find_keyword_position(sentence, keyword):
    p = []
    f = re.search(keyword, str(sentence))
    if f == None:
        return None
    else:
        while True:
            f = re.search(keyword, str(sentence))
            if f == None:
                break
            temp = re.search(keyword, str(sentence)).span()
            p.append(temp[1])

            sentence = sentence[temp[1]:]

        true_p = [p[0]]
        if len(p) > 1:  # 保证多个关键词，未测试
            for i in range(1, len(p)):
                true_p.append(true_p[i - 1] + p[i])
        return true_p


def find_keyword_context(sentence, keyword, true_p, number=2):
    keyword_len = len(keyword)
    sentence_len = len(sentence)
    result = []
    for i in range(len(true_p)):

        start = true_p[i] - keyword_len
        end = true_p[i]
        if start - number >= -1 and end + number <= sentence_len - 1:
            context = sentence[start - number:end + number]
            # print(context)
            result.append(context)
        elif start - number < -1 and end + number <= sentence_len - 1:
            context = sentence[0:end + number]
            # print(context)
            result.append(context)
        elif start - number >= -1 and end + number > sentence_len - 1:
            context = sentence[start - number:]
            # print(context)
            result.append(context)
        else:
            context = sentence
            # print(context)
            result.append(context)
    return result


def keyword4word(path, keyword, number):
    """
    有关键词返回关键词上下文列表
    否则换回空列表
    """
    total_result = []
    file = docx.Document(path)
    file_paragraphs_num = len(file.paragraphs)
    for i in range(file_paragraphs_num):
        sentence = file.paragraphs[i].text
        true_p = find_keyword_position(file.paragraphs[i].text, keyword)
        if true_p != None:
            result = find_keyword_context(file.paragraphs[i].text, keyword=keyword, true_p=true_p)
        else:
            continue
        total_result.append(result)
    return total_result


def get_file_keyword(keyword, number=2):
    # root = os.getcwd()
    fileList = get_file_list()
    file_dic = {}
    file_keyword = {}
    if keyword == "": return file_dic, 0
    for f in fileList:
        if f.endswith('.pdf'):
            pass
        elif f.endswith('.txt'):
            txt_result = find_keywords(f, keyword=keyword, number=number)
            if txt_result != []:
                file_dic[f] = txt_result
                file_keyword[f] = keyword
        elif f.endswith('.docx'):
            total_result = keyword4word(f, keyword=keyword, number=number)
            if total_result != []:
                file_dic[f] = total_result
                file_keyword[f] = keyword
        else:
            continue
    count_data = file_dic.values()
    count = 0
    for i in count_data:
        for j in i:
            for k in j:
                count += 1
    return file_dic, count
    # print(file_keyword)


if __name__ == "__main__":
    data = get_file_keyword("关键词")
    print(data[0].values(), data[1])
